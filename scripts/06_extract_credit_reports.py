#!/usr/bin/env python3
"""
06_extract_credit_reports.py
============================
Extracts text from existing credit papers in Word format (DOCX/DOC)
from the credit_reports/ folder and converts them into training pairs.

These are GOLD STANDARD training examples — real analyst-written credit
papers that teach the model exactly the output quality and format we want.

For each DOCX it:
  1. Extracts full text preserving paragraph structure
  2. Detects which CAMELS sections are present
  3. Splits into CAMELS pillar sections where possible
  4. Creates high-quality "assistant" training pairs from the real content

Run:
    python scripts/06_extract_credit_reports.py
    python scripts/06_extract_credit_reports.py --reprocess

Output:
    processed/credit_reports/<stem>.json
    training_data/credit_report_pairs.jsonl  (high-quality training pairs)
"""

import re
import json
import argparse
import traceback
from pathlib import Path
from datetime import datetime
from typing import Optional

PROJECT_ROOT       = Path(__file__).resolve().parent.parent
CREDIT_REPORTS_DIR = PROJECT_ROOT / "credit_reports"
OUTPUT_DIR         = PROJECT_ROOT / "processed" / "credit_reports"
TRAINING_DIR       = PROJECT_ROOT / "training_data"
LOGS_DIR           = PROJECT_ROOT / "logs"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
TRAINING_DIR.mkdir(exist_ok=True)
LOGS_DIR.mkdir(exist_ok=True)

SYSTEM_PROMPT = (
    "You are a senior credit analyst specialising in bank credit analysis "
    "using the CAMELS framework (Capital Adequacy, Asset Quality, Management, "
    "Earnings, Liquidity, Sensitivity to Market Risk). "
    "You follow Moody's, S&P Global Ratings, and Fitch Ratings methodologies. "
    "Every numerical claim must cite a source in the format [Source: p.XX] or "
    "[Source: Table Y, p.XX]. "
    "If data is unavailable, state 'Data not available' — never fabricate figures. "
    "Structure responses with: Assessment (Strong/Adequate/Weak/Critical), "
    "key metrics, peer context, risks, and source citations."
)

# Section header patterns — matches common credit paper headings
SECTION_PATTERNS = {
    "capital_adequacy": [
        r"capital\s+adequacy", r"capital\s+position", r"capitalisation",
        r"\bcapital\b.*\bcamels\b", r"C\s*[-–]\s*Capital",
    ],
    "asset_quality": [
        r"asset\s+quality", r"credit\s+quality", r"loan\s+quality",
        r"A\s*[-–]\s*Asset", r"non.performing", r"impairment",
    ],
    "management": [
        r"management\s+quality", r"governance", r"management\s+assessment",
        r"M\s*[-–]\s*Management",
    ],
    "earnings": [
        r"earnings", r"profitability", r"financial\s+performance",
        r"E\s*[-–]\s*Earnings", r"income\s+analysis",
    ],
    "liquidity": [
        r"liquidity", r"funding", r"L\s*[-–]\s*Liquidity",
        r"liquidity\s+(?:and|&)\s+funding",
    ],
    "sensitivity": [
        r"sensitivity", r"market\s+risk", r"S\s*[-–]\s*Sensitivity",
        r"interest\s+rate\s+risk",
    ],
    "overall_assessment": [
        r"overall\s+(?:assessment|rating|conclusion)",
        r"summary", r"conclusion", r"rating\s+rationale",
        r"executive\s+summary",
    ],
}

SECTION_TO_PILLAR = {
    "capital_adequacy":   "Capital Adequacy (C)",
    "asset_quality":      "Asset Quality (A)",
    "management":         "Management Quality (M)",
    "earnings":           "Earnings (E)",
    "liquidity":          "Liquidity & Funding (L)",
    "sensitivity":        "Sensitivity to Market Risk (S)",
    "overall_assessment": "Overall Assessment",
}


def check_docx_support() -> bool:
    try:
        import docx
        return True
    except ImportError:
        return False


def install_docx():
    import subprocess, sys
    subprocess.run(
        [sys.executable, "-m", "pip", "install", "python-docx", "--quiet"],
        check=True
    )


def extract_docx_text(docx_path: Path) -> tuple[str, list]:
    """
    Extract full text and paragraph list from a DOCX file.
    Returns (full_text, paragraphs_with_style).
    """
    from docx import Document
    doc = Document(str(docx_path))

    paragraphs = []
    for para in doc.paragraphs:
        text  = para.text.strip()
        style = para.style.name if para.style else "Normal"
        if text:
            paragraphs.append({"text": text, "style": style})

    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    paragraphs.append({"text": text, "style": "Table"})

    full_text = "\n\n".join(p["text"] for p in paragraphs)
    return full_text, paragraphs


def detect_section(text: str) -> Optional[str]:
    """Identify which CAMELS section a paragraph belongs to."""
    text_lower = text.lower()
    for section, patterns in SECTION_PATTERNS.items():
        for pat in patterns:
            if re.search(pat, text_lower):
                return section
    return None


def is_section_header(para: dict) -> bool:
    """Check if a paragraph is a section heading."""
    style = para.get("style", "Normal")
    text  = para["text"]
    # Heading styles
    if any(h in style for h in ["Heading", "Title", "Header"]):
        return True
    # Short uppercase or bold-like lines
    if len(text) < 80 and (text.isupper() or re.match(r"^[A-Z][\w\s\-&()]+$", text)):
        return True
    return False


def split_into_sections(paragraphs: list) -> dict:
    """
    Split document paragraphs into CAMELS sections.
    Returns dict of section_name -> text content.
    """
    sections = {}
    current_section = "introduction"
    current_content = []

    for para in paragraphs:
        text = para["text"]

        # Check if this paragraph is a section header
        if is_section_header(para):
            detected = detect_section(text)
            if detected:
                # Save current section
                if current_content:
                    sections[current_section] = "\n\n".join(current_content)
                current_section = detected
                current_content = [text]
                continue

        current_content.append(text)

    # Save final section
    if current_content:
        existing = sections.get(current_section, "")
        sections[current_section] = (existing + "\n\n" + "\n\n".join(current_content)).strip()

    return sections


def guess_bank_name(text: str, filename: str) -> str:
    bank_patterns = [
        r"(Lloyds\s+Banking\s+Group)", r"(HSBC\s+Holdings?)",
        r"(Barclays\s+(?:PLC|Bank)?)", r"(NatWest\s+Group)",
        r"(Standard\s+Chartered)", r"(JPMorgan\s+Chase)",
        r"(Bank\s+of\s+America)", r"(Wells\s+Fargo)",
        r"(Goldman\s+Sachs)", r"(Morgan\s+Stanley)",
        r"(Citigroup|Citi\s+Inc\.?)", r"(Deutsche\s+Bank)",
        r"(BNP\s+Paribas)", r"(UniCredit)",
    ]
    for pat in bank_patterns:
        m = re.search(pat, text[:2000], re.IGNORECASE)
        if m:
            return m.group(1).strip()
    stem = Path(filename).stem.replace("_", " ").replace("-", " ")
    return stem.title()


def guess_year(text: str, filename: str) -> Optional[str]:
    for pat in [r"(\d{4})\s+(?:Annual|Credit|Analysis)", r"(?:FY|fy)\s*(\d{4})",
                r"31\s+December\s+(\d{4})", r"December\s+31,?\s+(\d{4})"]:
        m = re.search(pat, text[:1000])
        if m:
            year = int(m.group(1))
            if 2000 <= year <= 2030:
                return str(year)
    m = re.search(r"20\d{2}", filename)
    if m:
        return m.group(0)
    return None


def build_training_pairs(
    bank_name: str,
    year: Optional[str],
    sections: dict,
    source_file: str,
) -> list:
    """
    Convert extracted sections into training pairs.
    These are GOLD STANDARD pairs — real analyst content.
    """
    pairs = []
    year_str = year or "unknown year"

    for section_key, content in sections.items():
        if section_key in ("introduction",) or len(content.split()) < 50:
            continue

        pillar_label = SECTION_TO_PILLAR.get(section_key,
                       section_key.replace("_", " ").title())

        user_content = (
            f"Bank: {bank_name}\n"
            f"Reporting Year: {year_str}\n\n"
            f"TASK: Analyse the {pillar_label} section for {bank_name}. "
            f"This is a CAMELS credit analysis. Provide a structured assessment "
            f"with key metrics, rating agency commentary, and key risks. "
            f"Cite specific data points with source references where available.\n\n"
            f"--- CONTEXT ---\n"
            f"Based on the {year_str} annual report and available financial data."
        )

        # The assistant response IS the real content from the credit paper
        assistant_content = (
            f"## {pillar_label}\n\n"
            f"**Bank:** {bank_name} | **Year:** {year_str}\n\n"
            f"{content}"
        )

        pairs.append({
            "messages": [
                {"role": "system",    "content": SYSTEM_PROMPT},
                {"role": "user",      "content": user_content},
                {"role": "assistant", "content": assistant_content},
            ],
            "_meta": {
                "source_file":  source_file,
                "bank_name":    bank_name,
                "year":         year_str,
                "section":      section_key,
                "pipeline":     "credit_report",
                "quality":      "gold",   # highest quality — real analyst content
            }
        })

    return pairs


def process_docx(docx_path: Path) -> Optional[dict]:
    print(f"\n  Processing: {docx_path.name}")

    try:
        full_text, paragraphs = extract_docx_text(docx_path)
    except Exception as e:
        print(f"    ❌ Failed to read DOCX: {e}")
        return None

    print(f"    Paragraphs: {len(paragraphs)}")
    print(f"    Characters: {len(full_text):,}")

    bank_name = guess_bank_name(full_text, docx_path.name)
    year      = guess_year(full_text, docx_path.name)
    sections  = split_into_sections(paragraphs)

    print(f"    Bank: {bank_name} | Year: {year or 'unknown'}")
    print(f"    Sections found: {[s for s in sections if s != 'introduction']}")

    training_pairs = build_training_pairs(bank_name, year, sections, docx_path.name)
    print(f"    Training pairs: {len(training_pairs)} (quality: GOLD)")

    return {
        "metadata": {
            "source_file":   docx_path.name,
            "source_path":   str(docx_path),
            "bank_name":     bank_name,
            "reporting_year": year,
            "total_paragraphs": len(paragraphs),
            "sections_found": list(sections.keys()),
            "training_pairs": len(training_pairs),
            "processed_at":  datetime.now().isoformat(),
        },
        "sections":       sections,
        "training_pairs": training_pairs,
    }


def write_jsonl(records: list, path: Path):
    with open(path, "w", encoding="utf-8") as f:
        for r in records:
            clean = {k: v for k, v in r.items() if not k.startswith("_")}
            f.write(json.dumps(clean, ensure_ascii=False) + "\n")
    print(f"  Written: {path.name} ({len(records)} records)")


def main():
    parser = argparse.ArgumentParser(
        description="Extract credit papers from Word DOCX files"
    )
    parser.add_argument("--reprocess", action="store_true")
    args = parser.parse_args()

    # Ensure folder exists
    CREDIT_REPORTS_DIR.mkdir(exist_ok=True)
    readme = CREDIT_REPORTS_DIR / "README.txt"
    if not readme.exists():
        readme.write_text(
            "Place existing credit papers (Word DOCX format) in this folder.\n"
            "These become GOLD STANDARD training examples.\n"
            "Run: ./run.sh --reprocess\n"
        )

    # Check for python-docx
    if not check_docx_support():
        print("Installing python-docx...")
        install_docx()

    # Find DOCX files
    docx_files = sorted(
        list(CREDIT_REPORTS_DIR.glob("*.docx")) +
        list(CREDIT_REPORTS_DIR.glob("*.doc"))
    )

    if not docx_files:
        print(f"\nNo DOCX files found in {CREDIT_REPORTS_DIR}")
        print(f"Place Word credit papers there to add gold-standard training data.")
        print(f"Pipeline will continue without them.")
        return

    print(f"\nFound {len(docx_files)} credit report(s) in credit_reports/")
    processed = skipped = errors = 0
    all_training_pairs = []

    for docx_path in docx_files:
        out_path = OUTPUT_DIR / (docx_path.stem + ".json")
        if out_path.exists() and not args.reprocess:
            print(f"  Skipping (exists): {docx_path.name}")
            skipped += 1
            # Load existing pairs for JSONL rebuild
            with open(out_path) as f:
                existing = json.load(f)
            all_training_pairs.extend(existing.get("training_pairs", []))
            continue

        try:
            result = process_docx(docx_path)
            if result is None:
                errors += 1
                continue
            with open(out_path, "w", encoding="utf-8") as f:
                json.dump(result, f, indent=2, ensure_ascii=False)
            all_training_pairs.extend(result.get("training_pairs", []))
            print(f"    ✅ {out_path.name}")
            processed += 1
        except Exception as e:
            print(f"    ❌ ERROR: {e}")
            traceback.print_exc()
            errors += 1

    # Write training pairs JSONL
    if all_training_pairs:
        pairs_path = TRAINING_DIR / "credit_report_pairs.jsonl"
        write_jsonl(all_training_pairs, pairs_path)
        print(f"\n✅ Gold-standard training pairs: {len(all_training_pairs)}")
        print(f"   These will be included in the next training run.")
    else:
        print(f"\nNo training pairs generated.")

    print(f"\nDone. Processed: {processed} | Skipped: {skipped} | Errors: {errors}")


if __name__ == "__main__":
    main()
