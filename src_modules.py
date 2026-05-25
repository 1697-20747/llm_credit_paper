#!/usr/bin/env python3
"""
src/pdf_extractor.py
====================
Extracts text and tables from annual report PDFs.
Produces a structured JSON with page-level text and tables,
each table cell referenced to its source page number.
"""

import re
import json
from pathlib import Path
from typing import Any
from loguru import logger

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import camelot
except ImportError:
    camelot = None


class PDFExtractor:
    """
    Extracts structured content from a bank annual report PDF.
    Prioritises PyMuPDF for text, pdfplumber for tables.
    Falls back to camelot for complex lattice tables.
    """

    def __init__(self, pdf_path: Path):
        self.pdf_path = Path(pdf_path)
        if not self.pdf_path.exists():
            raise FileNotFoundError(f"PDF not found: {pdf_path}")

    def extract(self) -> dict:
        logger.info(f"Extracting: {self.pdf_path.name}")
        pages = self._extract_with_pymupdf()
        pages = self._augment_tables_with_pdfplumber(pages)
        return {
            "source_file": str(self.pdf_path),
            "total_pages": len(pages),
            "pages": pages,
        }

    def _extract_with_pymupdf(self) -> list[dict]:
        if fitz is None:
            raise ImportError("PyMuPDF not installed: pip install PyMuPDF")
        doc = fitz.open(str(self.pdf_path))
        pages = []
        for page_num, page in enumerate(doc, start=1):
            text = page.get_text("text")
            blocks = page.get_text("blocks")  # (x0,y0,x1,y1, text, block_no, block_type)
            pages.append({
                "page_num": page_num,
                "text": text,
                "blocks": [
                    {"bbox": b[:4], "text": b[4], "type": b[6]}
                    for b in blocks
                ],
                "tables": [],  # Populated by pdfplumber pass
            })
        doc.close()
        logger.info(f"PyMuPDF extracted {len(pages)} pages")
        return pages

    def _augment_tables_with_pdfplumber(self, pages: list[dict]) -> list[dict]:
        if pdfplumber is None:
            logger.warning("pdfplumber not available — table extraction limited")
            return pages
        with pdfplumber.open(str(self.pdf_path)) as pdf:
            for page_data in pages:
                page_num = page_data["page_num"]
                if page_num > len(pdf.pages):
                    continue
                page = pdf.pages[page_num - 1]
                tables = page.extract_tables()
                if tables:
                    for tbl_idx, tbl in enumerate(tables):
                        # Infer caption from text above table
                        caption = self._infer_table_caption(
                            page_data["text"], tbl_idx
                        )
                        page_data["tables"].append({
                            "table_index": tbl_idx,
                            "page_num": page_num,
                            "caption": caption,
                            "data": tbl,
                            "source": f"p.{page_num}, Table {tbl_idx+1}",
                        })
        return pages

    def _infer_table_caption(self, page_text: str, table_idx: int) -> str:
        """Heuristically extract table caption from surrounding text."""
        # Look for lines like "Table 3: Consolidated Balance Sheet"
        patterns = [
            r"(?:Table|Exhibit|Figure|Note)\s+\d+[:\.\s]+([^\n]+)",
            r"((?:Consolidated|Group|Summary|Key)\s+[A-Z][^\n]+Sheet[^\n]*)",
            r"((?:Income|Balance|Capital|Risk|Liquidity|Funding)[^\n]{5,60})",
        ]
        for pat in patterns:
            matches = re.findall(pat, page_text)
            if matches:
                return matches[min(table_idx, len(matches)-1)].strip()
        return f"Table {table_idx + 1}"


# =============================================================================
# src/camels_mapper.py
# =============================================================================

"""
Maps raw PDF extraction to CAMELS financial taxonomy.
Every metric is stored with its source page reference.
"""

import re
from typing import Optional
from pydantic import BaseModel, Field


class SourcedMetric(BaseModel):
    value: Optional[float] = None
    raw_text: Optional[str] = None
    source_page: Optional[int] = None
    source_table: Optional[str] = None
    unit: Optional[str] = None   # "%", "£bn", "£m", "bps", "x"
    note: Optional[str] = None


class CapitalData(BaseModel):
    cet1_ratio: Optional[SourcedMetric] = None
    tier1_ratio: Optional[SourcedMetric] = None
    total_capital_ratio: Optional[SourcedMetric] = None
    leverage_ratio: Optional[SourcedMetric] = None
    rwa_bn: Optional[SourcedMetric] = None
    cet1_capital_bn: Optional[SourcedMetric] = None
    mrel_ratio: Optional[SourcedMetric] = None
    mrel_requirement: Optional[SourcedMetric] = None
    regulatory_min_cet1: Optional[SourcedMetric] = None
    rwa_density: Optional[SourcedMetric] = None


class AssetQualityData(BaseModel):
    total_gross_loans_bn: Optional[SourcedMetric] = None
    stage1_pct: Optional[SourcedMetric] = None
    stage2_pct: Optional[SourcedMetric] = None
    stage3_pct: Optional[SourcedMetric] = None
    stage1_bn: Optional[SourcedMetric] = None
    stage2_bn: Optional[SourcedMetric] = None
    stage3_bn: Optional[SourcedMetric] = None
    total_ecl_bn: Optional[SourcedMetric] = None
    stage3_ecl_bn: Optional[SourcedMetric] = None
    stage3_coverage_pct: Optional[SourcedMetric] = None
    total_coverage_pct: Optional[SourcedMetric] = None
    impairment_charge_bn: Optional[SourcedMetric] = None
    cost_of_risk_bps: Optional[SourcedMetric] = None
    net_writeoffs_bn: Optional[SourcedMetric] = None


class ManagementData(BaseModel):
    board_total: Optional[SourcedMetric] = None
    independent_neds: Optional[SourcedMetric] = None
    ceo_tenure_years: Optional[SourcedMetric] = None
    cro_reports_to: Optional[str] = None
    external_auditor: Optional[str] = None
    regulatory_actions_3yr: Optional[str] = None
    conduct_fines_5yr_m: Optional[SourcedMetric] = None
    ceo_total_comp_m: Optional[SourcedMetric] = None
    ceo_variable_pct: Optional[SourcedMetric] = None


class EarningsData(BaseModel):
    nii_bn: Optional[SourcedMetric] = None
    nim_pct: Optional[SourcedMetric] = None
    total_income_bn: Optional[SourcedMetric] = None
    non_interest_income_bn: Optional[SourcedMetric] = None
    operating_costs_bn: Optional[SourcedMetric] = None
    cost_income_ratio_pct: Optional[SourcedMetric] = None
    impairment_charge_bn: Optional[SourcedMetric] = None
    pretax_profit_bn: Optional[SourcedMetric] = None
    tax_charge_bn: Optional[SourcedMetric] = None
    pat_bn: Optional[SourcedMetric] = None
    rote_pct: Optional[SourcedMetric] = None
    roa_pct: Optional[SourcedMetric] = None
    eps_pence: Optional[SourcedMetric] = None
    dps_pence: Optional[SourcedMetric] = None
    prior_year_nim_pct: Optional[SourcedMetric] = None
    prior_year_rote_pct: Optional[SourcedMetric] = None


class LiquidityData(BaseModel):
    lcr_pct: Optional[SourcedMetric] = None
    nsfr_pct: Optional[SourcedMetric] = None
    hqla_bn: Optional[SourcedMetric] = None
    loan_deposit_ratio_pct: Optional[SourcedMetric] = None
    retail_deposits_bn: Optional[SourcedMetric] = None
    wholesale_funding_bn: Optional[SourcedMetric] = None
    wholesale_lt1yr_bn: Optional[SourcedMetric] = None
    wholesale_gt1yr_bn: Optional[SourcedMetric] = None
    liquidity_pool_bn: Optional[SourcedMetric] = None
    tfsme_bn: Optional[SourcedMetric] = None


class SensitivityData(BaseModel):
    nii_sensitivity_plus100bps_m: Optional[SourcedMetric] = None
    nii_sensitivity_minus100bps_m: Optional[SourcedMetric] = None
    nii_sensitivity_plus25bps_m: Optional[SourcedMetric] = None
    trading_var_99_1day_m: Optional[SourcedMetric] = None
    stressed_var_m: Optional[SourcedMetric] = None
    fx_net_open_bn: Optional[SourcedMetric] = None
    duration_of_equity_years: Optional[SourcedMetric] = None
    fvoci_portfolio_bn: Optional[SourcedMetric] = None
    fvoci_unrealised_loss_bn: Optional[SourcedMetric] = None
    irrbb_capital_bn: Optional[SourcedMetric] = None
    pension_surplus_deficit_m: Optional[SourcedMetric] = None


class CAMELSData(BaseModel):
    bank_name: str
    reporting_year: Optional[str] = None
    source_file: Optional[str] = None
    capital: CapitalData = Field(default_factory=CapitalData)
    asset_quality: AssetQualityData = Field(default_factory=AssetQualityData)
    management: ManagementData = Field(default_factory=ManagementData)
    earnings: EarningsData = Field(default_factory=EarningsData)
    liquidity: LiquidityData = Field(default_factory=LiquidityData)
    sensitivity: SensitivityData = Field(default_factory=SensitivityData)


class CAMELSMapper:
    """
    Maps raw PDF extraction dict to structured CAMELSData.
    Uses regex patterns calibrated for UK bank disclosures.
    """

    # Patterns: (field_name, regex, unit, page_hint_keywords)
    CAPITAL_PATTERNS = [
        ("cet1_ratio", r"CET\s*1\s+(?:ratio|capital ratio)[:\s]+(\d+\.?\d*)\s*%", "%",
         ["capital", "cet1"]),
        ("tier1_ratio", r"Tier\s*1\s+(?:ratio|capital ratio)[:\s]+(\d+\.?\d*)\s*%", "%",
         ["capital", "tier 1"]),
        ("total_capital_ratio", r"Total\s+capital\s+ratio[:\s]+(\d+\.?\d*)\s*%", "%",
         ["capital"]),
        ("leverage_ratio", r"(?:UK\s+)?[Ll]everage\s+ratio[:\s]+(\d+\.?\d*)\s*%", "%",
         ["leverage", "capital"]),
        ("rwa_bn", r"[Rr]isk[- ]?weighted\s+assets?[:\s£]+(\d[\d,\.]+)\s*(?:bn|billion)?", "£bn",
         ["rwa", "risk-weighted"]),
        ("mrel_ratio", r"MREL[:\s]+(\d+\.?\d*)\s*%", "%", ["mrel", "resolution"]),
    ]

    ASSET_QUALITY_PATTERNS = [
        ("stage1_pct", r"[Ss]tage\s*1[:\s]+(?:£[\d\.]+bn[,\s]+)?(\d+\.?\d*)\s*%", "%",
         ["stage", "ifrs", "ecl"]),
        ("stage2_pct", r"[Ss]tage\s*2[:\s]+(?:£[\d\.]+bn[,\s]+)?(\d+\.?\d*)\s*%", "%",
         ["stage", "ifrs", "ecl"]),
        ("stage3_pct", r"[Ss]tage\s*3[:\s]+(?:£[\d\.]+bn[,\s]+)?(\d+\.?\d*)\s*%", "%",
         ["stage", "ifrs", "ecl"]),
        ("cost_of_risk_bps", r"[Cc]ost\s+of\s+(?:credit\s+)?risk[:\s]+(\d+)\s*(?:bps|basis)", "bps",
         ["impairment", "credit risk"]),
        ("stage3_coverage_pct", r"[Ss]tage\s*3\s+coverage[:\s]+(\d+\.?\d*)\s*%", "%",
         ["coverage", "ecl"]),
    ]

    EARNINGS_PATTERNS = [
        ("nim_pct", r"[Nn]et\s+[Ii]nterest\s+[Mm]argin[:\s]+(\d+\.?\d*)\s*%", "%",
         ["net interest", "nim"]),
        ("cost_income_ratio_pct", r"[Cc]ost[:\s/]+[Ii]ncome\s+ratio[:\s]+(\d+\.?\d*)\s*%", "%",
         ["cost", "income", "efficiency"]),
        ("rote_pct", r"[Rr]eturn\s+on\s+[Tt]angible\s+[Ee]quity[:\s]+(\d+\.?\d*)\s*%", "%",
         ["rote", "return", "equity"]),
        ("roa_pct", r"[Rr]eturn\s+on\s+[Aa]ssets?[:\s]+(\d+\.?\d*)\s*%", "%",
         ["roa", "return on assets"]),
    ]

    LIQUIDITY_PATTERNS = [
        ("lcr_pct", r"[Ll]iquidity\s+[Cc]overage\s+[Rr]atio[:\s]+(\d+\.?\d*)\s*%", "%",
         ["lcr", "liquidity coverage"]),
        ("nsfr_pct", r"[Nn]et\s+[Ss]table\s+[Ff]unding\s+[Rr]atio[:\s]+(\d+\.?\d*)\s*%", "%",
         ["nsfr", "stable funding"]),
        ("loan_deposit_ratio_pct",
         r"[Ll]oan[s]?[:\s\-/]+[Dd]eposit[s]?\s+ratio[:\s]+(\d+\.?\d*)\s*%", "%",
         ["loan to deposit", "ldr"]),
    ]

    def __init__(self, extracted: dict):
        self.extracted = extracted
        self.pages = extracted.get("pages", [])

    def map(self) -> dict:
        """Returns CAMELS data as a plain dict (JSON-serialisable)."""
        full_text_by_page = {p["page_num"]: p["text"] for p in self.pages}

        capital = self._extract_section(self.CAPITAL_PATTERNS, full_text_by_page)
        asset_quality = self._extract_section(self.ASSET_QUALITY_PATTERNS, full_text_by_page)
        earnings = self._extract_section(self.EARNINGS_PATTERNS, full_text_by_page)
        liquidity = self._extract_section(self.LIQUIDITY_PATTERNS, full_text_by_page)

        # Tables extraction
        tables_summary = self._extract_key_tables()

        return {
            "bank_name": "Unknown",
            "source_file": self.extracted.get("source_file"),
            "total_pages": self.extracted.get("total_pages"),
            "capital": capital,
            "asset_quality": asset_quality,
            "management": {},
            "earnings": earnings,
            "liquidity": liquidity,
            "sensitivity": {},
            "tables_index": tables_summary,
            "extraction_notes": self._generate_notes(capital, asset_quality, earnings, liquidity),
        }

    def _extract_section(self, patterns: list, pages: dict) -> dict:
        results = {}
        full_text = "\n".join(pages.values())
        for field, pattern, unit, keywords in patterns:
            match = re.search(pattern, full_text)
            if match:
                raw_val = match.group(1).replace(",", "")
                try:
                    value = float(raw_val)
                except ValueError:
                    value = None
                # Find source page
                source_page = None
                for page_num, page_text in pages.items():
                    if re.search(pattern, page_text):
                        source_page = page_num
                        break
                results[field] = {
                    "value": value,
                    "raw_text": match.group(0)[:80],
                    "source_page": source_page,
                    "unit": unit,
                }
        return results

    def _extract_key_tables(self) -> list[dict]:
        tables = []
        for page in self.pages:
            for table in page.get("tables", []):
                caption = table.get("caption", "")
                # Flag financially relevant tables
                if any(kw in caption.lower() for kw in [
                    "balance", "income", "capital", "liquidity", "risk",
                    "funding", "asset", "ecl", "ifrs", "mrel"
                ]):
                    tables.append({
                        "caption": caption,
                        "page": table["page_num"],
                        "source": table["source"],
                        "rows": len(table.get("data", [])),
                    })
        return tables

    def _generate_notes(self, *sections) -> list[str]:
        notes = []
        for section in sections:
            missing = [k for k, v in section.items() if v is None]
            if missing:
                notes.append(f"Missing fields: {', '.join(missing)}")
        return notes


# =============================================================================
# src/llm_client.py
# =============================================================================

"""
Local LLM HTTP client.
Supports Ollama API (default) and llama.cpp server.
Zero external network calls.
"""

import asyncio
import json
from typing import Any
import httpx
from tenacity import retry, stop_after_attempt, wait_exponential
from loguru import logger


class LLMClient:
    """
    Async HTTP client for local LLM servers.
    Compatible with:
      - Ollama: http://localhost:11434
      - llama.cpp server: http://localhost:8080
    """

    def __init__(
        self,
        base_url: str = "http://localhost:11434",
        model: str = "camels-analyst",
        timeout: int = 300,
    ):
        self.base_url = base_url.rstrip("/")
        self.model = model
        self.timeout = timeout
        self._detect_backend()

    def _detect_backend(self):
        """Detect whether this is Ollama or llama.cpp."""
        if "11434" in self.base_url or "ollama" in self.base_url.lower():
            self.backend = "ollama"
        else:
            self.backend = "llamacpp"
        logger.debug(f"LLM backend: {self.backend} @ {self.base_url}")

    @retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=2, max=10))
    async def complete(
        self,
        messages: list[dict],
        temperature: float = 0.05,
        max_tokens: int = 2000,
    ) -> str:
        if self.backend == "ollama":
            return await self._ollama_complete(messages, temperature, max_tokens)
        else:
            return await self._llamacpp_complete(messages, temperature, max_tokens)

    async def _ollama_complete(self, messages, temperature, max_tokens) -> str:
        url = f"{self.base_url}/api/chat"
        payload = {
            "model": self.model,
            "messages": messages,
            "stream": False,
            "options": {
                "temperature": temperature,
                "num_predict": max_tokens,
                "top_p": 0.9,
                "repeat_penalty": 1.1,
            },
        }
        async with httpx.AsyncClient(timeout=self.timeout) as client:
            try:
                resp = await client.post(url, json=payload)
                resp.raise_for_status()
                data = resp.json()
                return data["message"]["content"]
            except httpx.ConnectError:
                raise ConnectionError(
                    f"Cannot connect to Ollama at {self.base_url}. "
                    f"Run: ollama serve"
                )

    async def _llamacpp_complete(self, messages, temperature, max_tokens) -> str:
        url = f"{self.base_url}/v1/chat/completions"
        payload = {
            "model": self.model,
            "messages": messages,
            "temperature": temperature,
            "max_tokens": max_tokens,
            "stream": False,
        }
        async with httpx.AsyncClient(timeout=self.timeout) as client:
            resp = await client.post(url, json=payload)
            resp.raise_for_status()
            data = resp.json()
            return data["choices"][0]["message"]["content"]

    async def health_check(self) -> bool:
        """Check if local LLM server is running."""
        try:
            async with httpx.AsyncClient(timeout=5) as client:
                if self.backend == "ollama":
                    resp = await client.get(f"{self.base_url}/api/tags")
                else:
                    resp = await client.get(f"{self.base_url}/health")
                return resp.status_code == 200
        except Exception:
            return False


# =============================================================================
# src/prompt_builder.py
# =============================================================================

"""
Builds structured prompts for CAMELS analysis.
Each prompt includes the raw extracted data so the LLM cannot fabricate.
"""


class PromptBuilder:

    system_prompt = (
        "You are a senior credit analyst at a major financial institution, "
        "specialising in bank credit analysis using the CAMELS framework. "
        "You follow rating agency methodologies: Moody's Banks methodology, "
        "S&P Global Ratings bank criteria, and Fitch Bank Rating Criteria. "
        "CRITICAL RULES:\n"
        "1. Every numerical claim MUST include [Source: p.XX] or [Source: Table Y].\n"
        "2. If data is not in the provided INPUT DATA, write 'Data not available' — "
        "   NEVER fabricate, estimate, or extrapolate.\n"
        "3. Use the exact numbers from the INPUT DATA — do not round or adjust.\n"
        "4. Begin each section with an Assessment: Strong / Adequate / Weak / Critical.\n"
        "5. End each section with Key Risks and a brief rating agency commentary.\n"
        "6. Flag any data inconsistencies you notice."
    )

    PILLAR_INSTRUCTIONS = {
        "capital": (
            "Analyse the Capital Adequacy (CAMELS — C) of the bank using the data below. "
            "Cover: CET1 ratio vs regulatory minimum, leverage ratio, MREL adequacy, "
            "RWA density considerations, Basel IV implications, and rating agency capital scoring."
        ),
        "asset_quality": (
            "Analyse Asset Quality (CAMELS — A) using the data below. "
            "Cover: IFRS 9 staging (Stage 1/2/3), ECL coverage adequacy, cost of risk, "
            "portfolio composition risks, and rating agency asset quality scoring."
        ),
        "management": (
            "Assess Management Quality and Governance (CAMELS — M) using the data below. "
            "Cover: board independence, risk governance structure, remuneration alignment, "
            "regulatory track record, audit quality, and strategic execution."
        ),
        "earnings": (
            "Analyse Earnings Quality and Sustainability (CAMELS — E) using the data below. "
            "Cover: RoTE vs cost of equity, NIM sustainability, cost efficiency, "
            "income diversification, and rating agency profitability scoring."
        ),
        "liquidity": (
            "Analyse Liquidity and Funding (CAMELS — L) using the data below. "
            "Cover: LCR/NSFR vs minimums, HQLA quality, funding mix stability, "
            "wholesale funding maturity profile, and structural liquidity risks."
        ),
        "sensitivity": (
            "Analyse Sensitivity to Market Risk (CAMELS — S) using the data below. "
            "Cover: interest rate sensitivity (NII impact), FVOCI unrealised losses, "
            "trading VaR, FX exposure, duration of equity, and pension risk."
        ),
    }

    def __init__(self, bank_name: str):
        self.bank_name = bank_name

    def build_pillar_prompt(self, pillar: str, data: dict) -> str:
        instruction = self.PILLAR_INSTRUCTIONS.get(pillar, "Analyse this CAMELS pillar.")
        data_str = self._format_data(data)
        return (
            f"Bank: {self.bank_name}\n\n"
            f"TASK: {instruction}\n\n"
            f"INPUT DATA (use ONLY these figures — cite source for each):\n"
            f"{data_str}\n\n"
            f"Write the analysis now. Begin with the Assessment rating."
        )

    def build_overall_rating_prompt(self, pillar_analyses: dict, camels_data: dict) -> str:
        summaries = "\n\n".join(
            f"=== {pillar.upper().replace('_', ' ')} ===\n{analysis[:500]}..."
            for pillar, analysis in pillar_analyses.items()
        )
        return (
            f"Bank: {self.bank_name}\n\n"
            f"Based on the following CAMELS pillar analyses, provide:\n"
            f"1. An overall CAMELS composite score (Strong / Adequate / Weak / Critical)\n"
            f"2. A composite rating equivalent (investment grade / sub-investment grade, "
            f"   with indicative rating range e.g. A-/Baa1/A-)\n"
            f"3. Key strengths (top 3)\n"
            f"4. Key vulnerabilities (top 3)\n"
            f"5. Rating outlook (Stable / Positive / Negative)\n"
            f"6. Recommended monitoring triggers\n\n"
            f"PILLAR SUMMARIES:\n{summaries}\n\n"
            f"Provide a structured overall assessment. "
            f"Do not introduce new figures not in the pillar analyses."
        )

    def _format_data(self, data: dict) -> str:
        if not data:
            return "(No data available for this pillar)"
        lines = []
        for key, metric in data.items():
            if isinstance(metric, dict):
                val = metric.get("value")
                unit = metric.get("unit", "")
                src = metric.get("source_page")
                src_str = f" [Source: p.{src}]" if src else " [Source: unknown page]"
                if val is not None:
                    lines.append(f"  {key.replace('_', ' ').title()}: {val}{unit}{src_str}")
                elif metric.get("raw_text"):
                    lines.append(f"  {key.replace('_', ' ').title()}: {metric['raw_text']}{src_str}")
            elif metric is not None:
                lines.append(f"  {key.replace('_', ' ').title()}: {metric}")
        return "\n".join(lines) if lines else "(No structured data extracted)"


# =============================================================================
# src/response_validator.py
# =============================================================================

"""
Validates LLM output against extracted source data.
Flags numbers in the analysis that don't appear in the CAMELS data.
"""

import re
from dataclasses import dataclass, field


@dataclass
class ValidationResult:
    has_hallucinations: bool = False
    flagged_items: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)


class ResponseValidator:
    """
    Extracts numbers from LLM response and checks they appear
    in the source CAMELS data. Flags discrepancies.
    """

    def __init__(self, camels_data: dict):
        self.camels_data = camels_data
        self.known_values = self._collect_known_values()

    def _collect_known_values(self) -> set[float]:
        values = set()
        for pillar in self.camels_data.values():
            if isinstance(pillar, dict):
                for metric in pillar.values():
                    if isinstance(metric, dict) and metric.get("value") is not None:
                        values.add(float(metric["value"]))
        return values

    def validate(self, pillar: str, analysis: str) -> ValidationResult:
        result = ValidationResult()

        # Extract all numbers from analysis
        numbers_in_analysis = re.findall(r"\b(\d+\.?\d*)\s*(?:%|bn|bps|x|m\b)", analysis)

        for num_str in numbers_in_analysis:
            try:
                num = float(num_str)
            except ValueError:
                continue
            # Check if this number (or close approximation) is in source data
            if num not in self.known_values:
                # Allow some tolerance for percentages and rounding
                if not any(abs(num - kv) < 0.1 for kv in self.known_values):
                    # Numbers under 10 are likely ratios/counts — don't flag
                    if num >= 10:
                        result.flagged_items.append(f"{num}")

        if result.flagged_items:
            result.has_hallucinations = True

        return result


# =============================================================================
# src/report_assembler.py
# =============================================================================

"""
Assembles the final credit paper from pillar analyses.
Produces Markdown, DOCX, and an audit index.
"""

import json
from datetime import datetime
from pathlib import Path


class ReportAssembler:

    PILLAR_DISPLAY = {
        "capital": "C — Capital Adequacy",
        "asset_quality": "A — Asset Quality",
        "management": "M — Management Quality",
        "earnings": "E — Earnings",
        "liquidity": "L — Liquidity & Funding",
        "sensitivity": "S — Sensitivity to Market Risk",
    }

    def __init__(self, bank_name, run_id, camels_data, pillar_analyses, overall_rating):
        self.bank_name = bank_name
        self.run_id = run_id
        self.camels_data = camels_data
        self.pillar_analyses = pillar_analyses
        self.overall_rating = overall_rating
        self.date = datetime.now().strftime("%d %B %Y")

    def to_markdown(self, path: Path):
        lines = [
            f"# Credit Analysis: {self.bank_name}",
            f"**Date:** {self.date}  ",
            f"**Methodology:** CAMELS Framework — aligned to Moody's / S&P / Fitch bank criteria  ",
            f"**Source:** Annual Report and Accounts  ",
            f"**Run ID:** `{self.run_id}`  ",
            "",
            "---",
            "",
            "## Overall Assessment",
            "",
            self.overall_rating,
            "",
            "---",
            "",
        ]
        for pillar_key, display_name in self.PILLAR_DISPLAY.items():
            analysis = self.pillar_analyses.get(pillar_key, "*Analysis not available*")
            lines += [
                f"## {display_name}",
                "",
                analysis,
                "",
                "---",
                "",
            ]
        lines += [
            "## Disclaimer",
            "",
            "This credit analysis was generated by an automated system using a locally-hosted "
            "AI model fine-tuned on credit analysis and rating agency methodology data. "
            "All figures are sourced from the bank's published annual report as cited. "
            "This document does not constitute investment advice. "
            "All analyses should be reviewed by a qualified credit professional before use.",
        ]
        path.write_text("\n".join(lines), encoding="utf-8")

    def to_docx(self, path: Path):
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            # If python-docx not available, just write markdown
            md_path = path.with_suffix(".md")
            self.to_markdown(md_path)
            return

        doc = Document()
        doc.add_heading(f"Credit Analysis: {self.bank_name}", 0)
        doc.add_paragraph(f"Date: {self.date}")
        doc.add_paragraph("Methodology: CAMELS Framework — Moody's / S&P / Fitch aligned")
        doc.add_paragraph(f"Run ID: {self.run_id}")
        doc.add_heading("Overall Assessment", level=1)
        doc.add_paragraph(self.overall_rating)

        for pillar_key, display_name in self.PILLAR_DISPLAY.items():
            doc.add_heading(display_name, level=1)
            analysis = self.pillar_analyses.get(pillar_key, "Analysis not available")
            doc.add_paragraph(analysis)

        doc.save(str(path))

    def to_audit_index(self, path: Path):
        """Builds an audit trail: every metric → source page."""
        index = {
            "bank_name": self.bank_name,
            "run_id": self.run_id,
            "generated_at": datetime.now().isoformat(),
            "source_citations": [],
        }
        for pillar, data in self.camels_data.items():
            if isinstance(data, dict):
                for metric, details in data.items():
                    if isinstance(details, dict) and details.get("source_page"):
                        index["source_citations"].append({
                            "pillar": pillar,
                            "metric": metric,
                            "value": details.get("value"),
                            "unit": details.get("unit"),
                            "source_page": details.get("source_page"),
                            "raw_text": details.get("raw_text", "")[:100],
                        })
        path.write_text(json.dumps(index, indent=2), encoding="utf-8")
